from __future__ import annotations

import os
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from PIL import Image
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Pillow is required in the bundled workspace Python runtime") from exc


ROOT = Path("/Users/xyq-mac/Documents/活动申请")
TEMPLATE = Path(
    "/Users/xyq-mac/Desktop/中关村数字经济产业联盟/案例征集/"
    "案例征集通知-2026全国企业数智化转型与赋能典型案例_副本.docx"
)
PPT = Path("/Users/xyq-mac/Desktop/中关村数字经济产业联盟/证元芳-产品介绍-20260529.pptx")
KB = Path("/private/tmp/Aicare-Product-Kb-activity")

OUT_DIR = ROOT / "outputs" / "2026全国企业数智化转型与赋能典型案例-证元芳"
MATERIALS_DIR = ROOT / "materials" / "2026全国企业数智化转型与赋能典型案例-证元芳"
IMAGES_DIR = OUT_DIR / "相关展示材料_JPG"


SOURCES = {
    "activity_notice": "案例征集通知-2026全国企业数智化转型与赋能典型案例_副本.docx",
    "kb_zyf": "qschouteam/Aicare-Product-Kb: references/06｜i-Magellanic 麦哲伦医药研究平台/证元芳.md",
    "kb_imagellanic": "qschouteam/Aicare-Product-Kb: references/06｜i-Magellanic 麦哲伦医药研究平台/平台总述.md",
    "kb_drgpt": "qschouteam/Aicare-Product-Kb: references/07｜大模型体系/Dr.GPT - Health OS.md",
    "kb_graph": "qschouteam/Aicare-Product-Kb: references/08｜业务场景模型矩阵/知识图谱.md",
    "kb_compliance": "qschouteam/Aicare-Product-Kb: references/97｜合规与监管资质.md",
    "kb_honor": "qschouteam/Aicare-Product-Kb: references/96｜行业荣誉与协会.md",
    "kb_core": "qschouteam/Aicare-Product-Kb: references/95｜集团对外核心数据卡.md",
    "product_ppt": "证元芳-产品介绍-20260529.pptx",
}


ENTERPRISE_INTRO = (
    "北京轻松怡康信息技术有限公司是轻松健康集团相关业务主体之一，围绕“健康+人工智能”双引擎战略，"
    "依托AIcare技术栈，推进健康服务、保险技术、医学研究辅助等场景的数智化应用。"
)

CASE_OVERVIEW = (
    "“证元芳”是面向医生、医学生及医学研究场景的循证医学智能体，以“问即有据”为核心。"
    "针对医学文献、指南、药品信息和病例知识分散，医生查证成本高，通用大模型依据不可追溯、难以进入严肃医疗工作流等痛点，"
    "产品以轻松问医Dr.GPT为底座，融合5000万+实体级别全球医学知识图谱，结合MedClaw多智能体协作和Agent DeepResearch，"
    "形成深度检索、权威性评估、准确性校验、引用验证的智能工作流，覆盖诊前查证、诊中决策参考、诊后患者解释，以及论文写作、"
    "临床研究方案、考试备考、医学图像、科普内容等场景。MedClaw Skills Store已上线2000+标准化Skill。"
    "2026年5月12日，证元芳通过中国信通院/泰尔实验室“医疗健康智能助手（MedClaw）能力测评”，13项功能用例全部通过，"
    "并已纳入北京市卫健委医疗AI应用评测体系（进行中），体现了医疗健康服务以专业大模型、知识图谱、智能体和合规测评驱动数智化转型的可复制路径。"
)

CASE_DETAIL_SHORT = (
    "详见附件一《证元芳：面向临床与医学科研的循证医学智能体案例详情》。附件围绕技术路径、实施模式、"
    "应用场景、实践成效、创新性、示范性与可推广性展开，并列明材料来源与待补充证明材料。"
)

EXHIBIT_TEXT = (
    "已整理产品介绍PPT中的产品首页、功能矩阵、DeepResearch流程、医学知识底座、测评/荣誉展示等JPG材料，"
    "详见“相关展示材料_JPG”文件夹；企业LOGO、主体建筑照片、盖章扫描件及正式联系人信息仍需申报单位补充。"
)


CASE_DETAIL_SECTIONS = [
    (
        "一、案例背景与痛点需求",
        [
            "医疗健康服务正在从经验驱动、人工检索驱动，走向以专业大模型、医学知识图谱和智能体工作流为核心的数智化服务模式。"
            "医生在临床接诊、患者解释、科研写作、临床研究方案设计和继续教育中，需要快速处理大量医学文献、指南、药品说明、"
            "真实世界研究资料与病例信息。传统方式依赖人工检索和个人经验，存在检索成本高、依据分散、更新滞后、不同医生之间知识调用不一致等问题。",
            "通用大模型虽然提升了问答效率，但在严肃医疗场景中仍面临依据不可追溯、回答可能存在幻觉、缺少医学专业评估闭环、"
            "难以满足医生对可验证证据的要求等风险。证元芳的建设目标不是替代医生诊疗决策，而是把AI能力嵌入医生的查证、解释、科研和学习工作流，"
            "为专业人士提供可核验、可追踪、可复用的循证支持。"
        ],
    ),
    (
        "二、建设目标与总体思路",
        [
            "本案例以“问即有据”为主线，面向临床医生、医学生、住院医师、医学运营和研究医生等专业用户，构建一个从循证问答起步，"
            "逐步扩展至医学科研、论文写作、临床研究方案、考试备考、科普内容生成、医学图像生成和日程管理的医学AI工作伙伴。",
            "总体思路是以轻松健康自研医疗-健康垂直大模型Dr.GPT提供医学语义理解、内容生成和推理能力，以5000万+实体级别全球医学知识融合底座提供结构化知识支撑，"
            "以MedClaw多智能体和Agent DeepResearch完成深度检索、权威性评估、准确性校验和引用验证，再通过移动App、微信公众号菜单栏、浏览器网页端等多入口融入医生日常工作。"
        ],
    ),
    (
        "三、技术路径",
        [
            "1. 专业大模型底座。证元芳上层能力基于轻松问医Dr.GPT构建。Dr.GPT是轻松健康自研医疗-健康垂直领域大模型，"
            "训练管线包含医学领域预训练、微调、基于可验证奖励的强化学习（RLVR）、人类指导强化学习、模型评估、医学评估和患者安全评估。"
            "知识库口径显示，Dr.GPT当前对外参数规模为100B，并有3000+医生参与人类指导强化学习训练。",
            "2. 医学知识融合底座。证元芳直接构建于全球医学知识融合底座之上，该底座融合PubMed、中文在线、北京卫和健康科技等数据源，"
            "形成5000万+实体级别医学知识图谱。其中卫和肿瘤专业结构化子集覆盖疾病、症状、检验检查、药品、临床路径、手术操作和医学知识图谱等维度，"
            "为肿瘤等高专业度场景提供结构化知识支撑。",
            "3. 多智能体协作与深度研究。证元芳通过MedClaw多智能体协作体系，将单轮医学问答扩展为多步骤研究式任务。"
            "Agent DeepResearch在问题拆解、检索策略、证据筛选、答案生成和引用验证之间形成工作流，确保回答不只给结论，还给出可追溯依据。",
            "4. 技能商店与平台化扩展。MedClaw Skills Store已上线2000+标准化Skill，覆盖临床诊疗、公共卫生、医学教育等多领域，"
            "支持第三方技能接入和一键调用，使证元芳从单一问答工具演进为开放的医学智能体能力平台。",
            "5. 安全合规与数据隔离。产品能力设计中包含企业级数据隔离、多端接入、用户记忆等能力。相关AI能力已纳入集团多项算法备案、"
            "北京网信办AI产品登记和医疗领域人工智能应用评测体系；在对外表述中，应区分证元芳自身测评、上游模型备案和集团整体资质，避免将集团资质直接等同于证元芳单产品备案。"
        ],
    ),
    (
        "四、实施模式",
        [
            "证元芳采用“医生免费 + 机构付费”的分层模式，降低医生个人使用门槛，将商业化重点放在医院、药企及相关机构的定制化付费、系统融合、培训部署和学术推广合作上。"
            "该模式有利于让AI工具先进入医生个人工作流，再通过机构级部署与院内系统、药企学术推广和医学研究项目结合。",
            "从产品阶段看，证元芳1.0阶段聚焦循证问答和证据来源追溯；当前2.0阶段已经扩展至医学生和更广泛医师群体，能力覆盖论文写作与优化、临床研究方案设计、考试备考、"
            "医学图像生成、科普内容生成、日程管理等。原“科研智能体蜂群”中的文献解读、临床研究方案撰写等能力，也已并入证元芳2.0科研工作流模块。"
        ],
    ),
    (
        "五、应用场景",
        [
            "1. 诊前：医生在接诊前快速获取与患者病情相关的文献、指南和循证依据，降低人工检索成本。"
            "2. 诊中：医生在治疗路径、用药方案或检查安排等环节获得可追溯的医学依据，辅助形成更稳健的专业判断。"
            "3. 诊后：医生可将专业证据转化为患者更容易理解的解释和随访建议，提升医患沟通效率。",
            "4. 医学科研：支持文献检索、文献解读、研究问题拆解、论文写作与优化、临床研究方案撰写等任务。"
            "5. 医学教育与培训：面向医学生、住院医师和专科医生提供真题备考、解析、知识点回顾和病例学习。"
            "6. 患者教育与科普：帮助医生生成通俗易懂的科普内容、患者教育材料和医学示意图。"
        ],
    ),
    (
        "六、实践成效",
        [
            "第一，权威测评验证。证元芳V1.0已通过中国信通院/泰尔实验室“医疗健康智能助手（MedClaw）能力测评”，检测日期为2026年4月29日至5月6日，"
            "颁发日期为2026年5月12日，证书编号为2026TZ000321，报告编号为26B01Z100580-003。测评共13项功能用例，实测13项，通过13项，未通过0项，"
            "覆盖循证问答即时响应、证据来源追溯、复杂病例深度分析、多Agent协作、企业级数据隔离、医学图像生成、科普内容生成、论文生成与优化、临床试验方案设计、"
            "技能商店调用、真题备考模式、多端接入和用户记忆功能。",
            "第二，专业能力验证。知识库显示，证元芳在CMB中国执业医师资格考试中实现国内首家满分，并在肿瘤科正高、副高考试中取得SOTA成绩。"
            "这些结果体现了产品在医学知识理解、专业推理和考试型标准化任务中的能力水平。",
            "第三，平台化与生态基础。证元芳所属的i-Magellanic麦哲伦医药研究平台截至2025年12月31日已落地46项数字医学研究辅助项目，合作103家制药合作伙伴；"
            "集团数字医学研究辅助服务收入2025年为41.7百万元，同比增长29.5%。这些平台层成果为证元芳在医学研究、药企合作和机构服务场景中的推广奠定基础。",
            "第四，合规与监管进程。证元芳已于2025年11月19日提交北京市卫生健康委员会医疗领域人工智能应用评测，目前处于评测进行中状态；"
            "对外材料中应表述为“已纳入评测体系”，不得写成“已通过”。"
        ],
    ),
    (
        "七、创新性、示范性与可推广性",
        [
            "创新性方面，证元芳将医疗-健康垂直大模型、全球医学知识图谱、Agent DeepResearch、多智能体协作和技能商店整合为统一医学智能体，"
            "把传统“搜索资料—人工判断—写作输出”的长流程压缩为可验证、可追溯的智能工作流。",
            "示范性方面，案例展示了医疗健康企业如何在合规边界内推进AI应用：既强调大模型能力，又通过第三方测评、引用验证、数据隔离和监管评测进程建立可信机制；"
            "既提升医生效率，又不替代医生最终诊疗决策，符合严肃医疗场景的责任边界。",
            "可推广性方面，证元芳的能力可在医院科室、药企医学事务、真实世界研究、医学教育、患者教育和健康科普等场景复用。"
            "其“专业大模型 + 领域知识图谱 + 智能体工作流 + 可追溯证据”的技术路径，也可为其他企业数智化转型提供参考。"
        ],
    ),
    (
        "八、真实性与材料边界说明",
        [
            "本申报材料中的产品能力、技术架构、测评结论、平台数据和合规状态来自Aicare产品知识库、活动通知及用户提供的证元芳产品介绍资料。"
            "当前知识库明确提示：证元芳暂无可公开披露的具体医院/科室落地案例，CMB满分、肿瘤科正高/副高SOTA属于能力验证，不应写成客户落地成效。",
            "为增强申报竞争力，建议申报单位补充至少一个真实使用案例，包括医院/科室或医生群体、上线时间、使用规模、典型场景、效率提升或质量提升数据、用户反馈、截图或证明文件。"
        ],
    ),
]


SOURCE_NOTES = [
    ("活动规则", "征集对象为企业数智化转型实践；材料含企业简介、案例概述、案例详情和JPG展示材料；截止日期为2026年5月29日。", SOURCES["activity_notice"]),
    ("申报对象", "申报单位：北京轻松怡康信息技术有限公司；申报案例：证元芳。", "用户指令"),
    ("产品定位", "证元芳是i-Magellanic麦哲伦医药研究平台下的医学AI助手，核心表达为“问即有据”。", SOURCES["kb_zyf"]),
    ("能力阶段", "证元芳当前处于2.0阶段，从循证扩展到论文、临床研究方案、备考、医学图像、科普、日程管理等场景。", SOURCES["kb_zyf"]),
    ("知识底座", "全球医学知识融合底座为5000万+实体级别，融合PubMed、中文在线、北京卫和健康科技等数据源。", SOURCES["kb_zyf"] + "；" + SOURCES["kb_graph"]),
    ("技能商店", "MedClaw Skills Store当前已上线2000+标准化Skill，覆盖临床诊疗、公共卫生、医学教育等领域。", SOURCES["kb_zyf"]),
    ("测评结果", "2026-05-12通过中国信通院/泰尔实验室医疗健康智能助手（MedClaw）能力测评，13/13用例通过，证书编号2026TZ000321。", SOURCES["kb_compliance"]),
    ("监管进程", "证元芳已提交北京市卫生健康委员会医疗领域人工智能应用评测，状态为进行中，不能写为已通过。", SOURCES["kb_compliance"]),
    ("平台成果", "i-Magellanic平台截至2025-12-31已落地46项数字医学研究辅助项目，合作103家制药合作伙伴。", SOURCES["kb_imagellanic"]),
    ("主体风险", "CAICT测评委托单位为北京轻松健康网络科技有限公司；如申报单位为北京轻松怡康信息技术有限公司，建议补充集团内授权或联合申报说明。", SOURCES["kb_compliance"]),
]


DISPLAY_IMAGES = {
    "image18.jpeg": "01_企业及品牌展示.jpg",
    "image5.png": "02_证元芳产品首页.jpg",
    "image6.png": "03_证元芳产品特性.jpg",
    "image7.png": "04_诊前诊中诊后循证支持.jpg",
    "image8.png": "05_证元芳全场景能力矩阵.jpg",
    "image14.png": "06_Agent_DeepResearch流程.jpg",
    "image15.png": "07_医学知识融合底座.jpg",
    "image16.png": "08_测评证书与荣誉展示.jpg",
    "image17.png": "09_行业协会与联盟身份.jpg",
    "image13.png": "10_能力矩阵与系统流程.jpg",
}


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_text(paragraph, text: str, size: float = 10.5, bold: bool = False) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.clear()


def fill_cell(cell, text: str, size: float = 10.5, bold: bool = False) -> None:
    clear_cell(cell)
    parts = text.split("\n")
    for idx, part in enumerate(parts):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        set_paragraph_text(p, part, size=size, bold=bold)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, font_size, color in [
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12.5, "1F4E79"),
        ("Heading 3", 11.5, "365F91"),
    ]:
        style = styles[name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def build_filled_form() -> Path:
    doc = Document(str(TEMPLATE))
    table = doc.tables[0]

    # Preserve the official table while filling only the response cells.
    fill_cell(table.rows[1].cells[1], "北京轻松怡康信息技术有限公司")
    fill_cell(table.rows[1].cells[3], "医疗健康服务 / 人工智能 / 数字经济")
    fill_cell(table.rows[2].cells[1], "待补充：请以营业执照注册地址为准")
    fill_cell(table.rows[3].cells[1], "待补充")
    fill_cell(table.rows[3].cells[3], "待补充")
    fill_cell(
        table.rows[5].cells[1],
        "□2026全国企业数智化转型与赋能十佳案例\n■2026全国企业数智化转型与赋能优秀案例",
    )
    fill_cell(table.rows[6].cells[1], "证元芳：面向临床与医学科研的循证医学智能体")
    fill_cell(table.rows[7].cells[1], "待补充：不超过3人，建议填写产品负责人、技术负责人、医学负责人")
    fill_cell(table.rows[8].cells[1], "待补充：不超过10人，建议覆盖产品、算法、医学、数据、合规、市场等角色")
    fill_cell(table.rows[9].cells[1], ENTERPRISE_INTRO, size=10)
    fill_cell(table.rows[10].cells[1], CASE_OVERVIEW + "\n\n（详见附件一）", size=9.5)
    fill_cell(table.rows[11].cells[1], CASE_DETAIL_SHORT, size=10)
    fill_cell(table.rows[12].cells[1], EXHIBIT_TEXT, size=10)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=run.font.size.pt if run.font.size else None)

    out = OUT_DIR / "01_2026全国企业数智化转型与赋能典型案例征集表_北京轻松怡康_证元芳.docx"
    doc.save(out)
    return out


def add_source_table(doc: Document) -> None:
    doc.add_heading("附录：主要事实来源", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "事实类型"
    hdr[1].text = "申报材料采用口径"
    hdr[2].text = "来源"
    for cell in hdr:
        shade_cell(cell, "D9EAF7")
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=9.5, bold=True)
    for item, fact, source in SOURCE_NOTES:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = fact
        cells[2].text = source
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_font(r, size=8.5)


def build_case_detail_doc() -> Path:
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026全国企业数智化转型与赋能典型案例申报材料")
    set_run_font(run, size=18, bold=True, color="1F4E79")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("证元芳：面向临床与医学科研的循证医学智能体")
    set_run_font(run, size=14, bold=True, color="365F91")

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("申报单位", "北京轻松怡康信息技术有限公司"),
        ("申报类别", "2026全国企业数智化转型与赋能优秀案例（如补充真实落地数据，可升级申报十佳案例）"),
        ("案例名称", "证元芳"),
        ("资料版本", "申报草案 v1.0，生成日期：2026年6月1日"),
    ]
    for idx, (k, v) in enumerate(rows):
        meta.rows[idx].cells[0].text = k
        meta.rows[idx].cells[1].text = v
        shade_cell(meta.rows[idx].cells[0], "EAF3F8")
        for cell in meta.rows[idx].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9.5, bold=(cell is meta.rows[idx].cells[0]))

    doc.add_heading("企业简介", level=1)
    doc.add_paragraph(ENTERPRISE_INTRO)

    doc.add_heading("案例概述", level=1)
    doc.add_paragraph(CASE_OVERVIEW)

    for heading, paragraphs in CASE_DETAIL_SECTIONS:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

    add_source_table(doc)

    out = OUT_DIR / "02_案例详情_证元芳_面向临床与医学科研的循证医学智能体.docx"
    doc.save(out)
    return out


def build_submission_notes() -> Path:
    content = f"""# 提交说明与待补充清单

## 活动规则摘要

- 活动名称：2026全国企业数智化转型与赋能典型案例
- 组织方：中关村数字经济产业联盟 / 中数联盟
- 征集对象：在技术创新、场景应用、提质增效、生态赋能等方面取得显著成效的企业项目
- 提交材料：Word文字材料 + JPG展示材料，打包压缩后发送至 `599321147@qq.com`
- 邮件主题格式：`2026全国企业数智化转型与赋能典型案例-单位简称-案例名称`
- 本次建议邮件主题：`2026全国企业数智化转型与赋能典型案例-轻松怡康-证元芳`
- 截止日期风险：通知写明截止日期为 `2026年5月29日`，当前生成日期为 `2026年6月1日`，存在逾期不予受理风险。

## 已生成材料

- `01_2026全国企业数智化转型与赋能典型案例征集表_北京轻松怡康_证元芳.docx`
- `02_案例详情_证元芳_面向临床与医学科研的循证医学智能体.docx`
- `相关展示材料_JPG/`：由用户提供的证元芳产品介绍PPT提取并转为JPG
- `product-kb-notes.md`：产品知识库事实溯源摘要

## 需要申报单位补充或确认

1. 申报单位注册地址：请以营业执照注册地址为准。
2. 联系人及联系方式：姓名、手机、邮箱。
3. 主创人及职务：不超过3人。
4. 参创人及职务：不超过10人。
5. 申报类别：当前按“优秀案例”填写；如希望冲刺“十佳案例”，建议先补充真实落地规模和量化成效。
6. 真实落地案例：至少补充1个医院/科室/医生群体使用证元芳的案例，包含上线时间、使用规模、典型场景、效率提升或质量提升数据、用户反馈与截图证明。
7. 主体归属说明：CAICT/泰尔实验室测评委托单位为“北京轻松健康网络科技有限公司”，本次申报单位为“北京轻松怡康信息技术有限公司”，建议补充集团内授权、联合申报说明或产品归属说明。
8. JPG展示材料：建议补充企业LOGO高清图、主体建筑照片、产品界面高清截图、数智化应用场景截图。
9. 原创声明：正式提交前需申报人签字、单位盖章，并确认日期填写策略。

## 事实边界提醒

- 可以写：证元芳通过中国信通院/泰尔实验室MedClaw能力测评，13/13用例通过。
- 可以写：证元芳已纳入北京市卫生健康委员会医疗领域人工智能应用评测体系，评测进行中。
- 不要写：证元芳已通过北京市卫健委评测。
- 不要写：证元芳本身已通过6项国家网信办算法备案。知识库明确提示，6项算法备案属于集团/关联公司整体，不等同于证元芳单产品备案。
- 不要把能力验证写成客户落地成效。CMB满分、肿瘤科正高/副高SOTA是能力验证；真实医院/科室案例仍待补充。
"""
    out = OUT_DIR / "03_提交说明与待补充清单.md"
    out.write_text(content, encoding="utf-8")
    return out


def build_kb_notes() -> Path:
    MATERIALS_DIR.mkdir(parents=True, exist_ok=True)
    lines = [
        "# 证元芳申报材料产品知识库检索摘要",
        "",
        "- 知识库：`qschouteam/Aicare-Product-Kb`",
        "- 检索方式：git 浅克隆至 `/private/tmp/Aicare-Product-Kb-activity` 后用 `rg` 深检索 `证元芳 / Dr.GPT / i-Magellanic / MedClaw / 知识图谱 / 合规`。",
        "- 当前远端 master：`c2c76bfbe0c79ebe6be8e68103e07156a5e9a036`",
        "- 申报对象：证元芳",
        "",
        "## 可写入申报材料的事实",
        "",
    ]
    for item, fact, source in SOURCE_NOTES:
        lines.append(f"### {item}")
        lines.append("")
        lines.append(f"- 事实：{fact}")
        lines.append(f"- 来源：{source}")
        lines.append("")
    lines.extend(
        [
            "## 不应写成确定事实的内容",
            "",
            "- 不应写“证元芳已通过6项国家网信办算法备案”，知识库说明6项备案属于集团/关联公司整体。",
            "- 不应写“证元芳已通过北京市卫健委评测”，当前状态为评测进行中。",
            "- 不应写具体医院/科室上线规模、活跃医生数、效率提升百分比，除非用户补充证明材料。",
        ]
    )
    out = MATERIALS_DIR / "product-kb-notes.md"
    out.write_text("\n".join(lines), encoding="utf-8")
    shutil.copy2(out, OUT_DIR / "04_product-kb-notes.md")
    return out


def extract_display_images() -> list[Path]:
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    tmp = OUT_DIR / "_ppt_media"
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir(parents=True)
    with zipfile.ZipFile(PPT) as zf:
        for name in zf.namelist():
            base = Path(name).name
            if name.startswith("ppt/media/") and base in DISPLAY_IMAGES:
                zf.extract(name, tmp)
                extracted = tmp / name
                target = IMAGES_DIR / DISPLAY_IMAGES[base]
                with Image.open(extracted) as im:
                    if im.mode in ("RGBA", "LA"):
                        bg = Image.new("RGB", im.size, "white")
                        alpha = im.getchannel("A")
                        bg.paste(im.convert("RGB"), mask=alpha)
                        out_im = bg
                    else:
                        out_im = im.convert("RGB")
                    out_im.save(target, "JPEG", quality=92, optimize=True)
    shutil.rmtree(tmp)
    return sorted(IMAGES_DIR.glob("*.jpg"))


def build_manifest(files: list[Path], images: list[Path]) -> Path:
    content = ["# 申报资料包目录", ""]
    content.append("## Word / 文档材料")
    for file in files:
        content.append(f"- `{file.relative_to(OUT_DIR)}`")
    content.append("")
    content.append("## JPG展示材料")
    for image in images:
        content.append(f"- `{image.relative_to(OUT_DIR)}`")
    content.append("")
    content.append("## 建议压缩包命名")
    content.append("")
    content.append("`2026全国企业数智化转型与赋能典型案例-轻松怡康-证元芳.zip`")
    out = OUT_DIR / "00_申报资料包目录.md"
    out.write_text("\n".join(content), encoding="utf-8")
    return out


def build_zip() -> Path:
    zip_path = ROOT / "outputs" / "2026全国企业数智化转型与赋能典型案例-轻松怡康-证元芳.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in OUT_DIR.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(OUT_DIR.parent))
    return zip_path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MATERIALS_DIR.mkdir(parents=True, exist_ok=True)

    form = build_filled_form()
    detail = build_case_detail_doc()
    notes = build_submission_notes()
    kb_notes = build_kb_notes()
    images = extract_display_images()
    manifest = build_manifest([form, detail, notes, OUT_DIR / "04_product-kb-notes.md"], images)
    zip_path = build_zip()

    print("Generated:")
    for path in [manifest, form, detail, notes, kb_notes, zip_path, *images]:
        print(path)


if __name__ == "__main__":
    main()
